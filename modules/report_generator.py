# modules/report_generator.py

import streamlit as st
from docx import Document
from io import BytesIO
import os

# ✅ 新版 openai
import openai
from openai import OpenAI, APIConnectionError, APIError, AuthenticationError

# ✅ PDF轉檔
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

def run():
    st.subheader("AI報告產出 + PDF + ChatGPT商業模式生成")

    if "answers" not in st.session_state:
        st.warning("請先填寫問卷資料。")
        return

    answers = st.session_state["answers"]
    bm_description = answers.get("bm_description", "")

    # ✅ 自動補推薦元素（防止 Day7 沒執行導致資料不一致）
    if "recommended_elements" not in st.session_state or not st.session_state["recommended_elements"]:
        import pandas as pd
        import jieba
        try:
            description = answers.get("bm_description", "")
            words = set(jieba.lcut(description))
            df = pd.read_excel("data/element_info.xlsx")
            recommended_elements = []
            for _, row in df.iterrows():
                element = row["元素"]
                keywords = row["關鍵詞"].split(",")
                if any(keyword.strip() in words for keyword in keywords):
                    recommended_elements.append(element)
            st.session_state["recommended_elements"] = recommended_elements
            st.info("✅ 自動補全推薦元素成功")
        except Exception as e:
            st.warning(f"❗ 自動計算推薦元素時發生錯誤：{e}")

    recommended_elements = st.session_state.get("recommended_elements", [])

    st.markdown("### 問卷資料")
    st.json(answers)

    st.markdown("### 推薦元素")
    st.write(recommended_elements)

    # ✅ AI自動寫商業模式
    st.markdown("### AI輔助生成商業模式")
    if st.button("👉 根據推薦元素請 ChatGPT 寫商業模式"):
        if "openai_api_key" not in st.secrets or not st.secrets["openai_api_key"]:
            st.error("❌ 未設定 OpenAI API Key，請在 .streamlit/secrets.toml 中加入 openai_api_key。")
        else:
            prompt = f"請根據以下關鍵字，幫我寫一段約300字的公司商業模式描述：{', '.join(recommended_elements)}"
            try:
                client = OpenAI(api_key=st.secrets["openai_api_key"])
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": prompt}]
                )
                generated_text = response.choices[0].message.content
                st.session_state["ai_business_text"] = generated_text
                st.success("✅ AI 商業模式已產生！")
                st.text_area("AI生成商業模式", value=generated_text, height=200)
            except (APIConnectionError, APIError, AuthenticationError) as e:
                st.error(f"❌ OpenAI API 發生錯誤：{e}")

    # ✅ 產出報告
    if st.button("📄 產出 Word + PDF 報告"):
        doc = Document()
        doc.add_heading('EcoStrategy AI 智慧問卷分析報告', level=1)

        doc.add_heading('一、企業基本資料', level=2)
        for key, value in answers.items():
            if key != "bm_description":
                doc.add_paragraph(f"{key}: {value}")

        doc.add_heading('二、商業模式敘述', level=2)
        doc.add_paragraph(bm_description)

        doc.add_heading('三、推薦相關元素', level=2)
        if recommended_elements:
            doc.add_paragraph("、".join(recommended_elements))
        else:
            doc.add_paragraph("未發現推薦元素。")

        doc.add_heading('四、AI輔助商業模式建議', level=2)
        ai_text = st.session_state.get("ai_business_text", "尚未產生AI建議。")
        doc.add_paragraph(ai_text)

        word_path = "temp_report.docx"
        doc.save(word_path)

        # ✅ PDF轉換
        if DOCX2PDF_AVAILABLE:
            pdf_path = "temp_report.pdf"
            try:
                convert(word_path, pdf_path)
                st.success("✅ PDF 報告已同步產生！")
            except Exception as e:
                st.warning(f"❗ 轉換 PDF 發生錯誤：{e}")
        else:
            pdf_path = None
            st.info("❗ 未安裝 docx2pdf，將只產出 Word 報告。")

        # ✅ 提供下載
        with open(word_path, "rb") as file:
            st.download_button(
                label="📥 下載 Word 報告",
                data=file,
                file_name="EcoStrategy_AI_分析報告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        if pdf_path and os.path.exists(pdf_path):
            with open(pdf_path, "rb") as file:
                st.download_button(
                    label="📥 下載 PDF 報告",
                    data=file,
                    file_name="EcoStrategy_AI_分析報告.pdf",
                    mime="application/pdf"
                )

        # ✅ 清除暫存
        if os.path.exists(word_path):
            os.remove(word_path)
        if pdf_path and os.path.exists(pdf_path):
            os.remove(pdf_path)